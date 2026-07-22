from __future__ import annotations

import json
import math
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "work"
FIG = ROOT / "figures"
OUT = ROOT / "RoadGen3D_技术交底书_当前LLM参数与2D-3D交互_2026-07-20.docx"
FIG.mkdir(parents=True, exist_ok=True)


def load(name: str):
    return json.loads((WORK / name).read_text(encoding="utf-8"))


SOURCES = load("01-source-map.json")
INVENTORY = load("02-technical-inventory.json")
EVIDENCE = load("03-evidence-ledger.json")
CLAIMS = load("04-claim-strategy.json")


FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
FONT_REG = ImageFont.truetype(FONT_PATH, 34)
FONT_SMALL = ImageFont.truetype(FONT_PATH, 27)
FONT_TINY = ImageFont.truetype(FONT_PATH, 23)
FONT_TITLE = ImageFont.truetype(FONT_PATH, 45)
FONT_BOLD = ImageFont.truetype(FONT_PATH, 32)


def _text_size(draw, text, font, spacing=8):
    b = draw.multiline_textbbox((0, 0), text, font=font, spacing=spacing, align="center")
    return b[2] - b[0], b[3] - b[1]


def _draw_centered(draw, rect, text, font=FONT_SMALL, spacing=8):
    x1, y1, x2, y2 = rect
    tw, th = _text_size(draw, text, font, spacing)
    draw.multiline_text(((x1 + x2 - tw) / 2, (y1 + y2 - th) / 2), text,
                        fill="black", font=font, spacing=spacing, align="center")


def _dash_rect(draw, rect, width=4, dash=18):
    x1, y1, x2, y2 = rect
    for x in range(x1, x2, dash * 2):
        draw.line((x, y1, min(x + dash, x2), y1), fill="black", width=width)
        draw.line((x, y2, min(x + dash, x2), y2), fill="black", width=width)
    for y in range(y1, y2, dash * 2):
        draw.line((x1, y, x1, min(y + dash, y2)), fill="black", width=width)
        draw.line((x2, y, x2, min(y + dash, y2)), fill="black", width=width)


def _box(draw, rect, text, title=None, dashed=False, width=4, font=FONT_SMALL):
    if dashed:
        _dash_rect(draw, rect, width)
    else:
        draw.rectangle(rect, outline="black", width=width, fill="white")
    x1, y1, x2, y2 = rect
    if title:
        draw.text((x1 + 16, y1 + 8), title, fill="black", font=FONT_BOLD)
        _draw_centered(draw, (x1 + 8, y1 + 42, x2 - 8, y2 - 6), text, font)
    else:
        _draw_centered(draw, rect, text, font)


def _arrow(draw, start, end, dashed=False, width=5):
    x1, y1 = start
    x2, y2 = end
    if dashed:
        segments = 12
        for i in range(segments):
            if i % 2 == 0:
                a = i / segments
                b = (i + 1) / segments
                draw.line((x1 + (x2-x1)*a, y1 + (y2-y1)*a, x1 + (x2-x1)*b, y1 + (y2-y1)*b), fill="black", width=width)
    else:
        draw.line((x1, y1, x2, y2), fill="black", width=width)
    angle = math.atan2(y2-y1, x2-x1)
    size = 22
    p1 = (x2 - size * math.cos(angle - 0.55), y2 - size * math.sin(angle - 0.55))
    p2 = (x2 - size * math.cos(angle + 0.55), y2 - size * math.sin(angle + 0.55))
    draw.polygon([(x2, y2), p1, p2], fill="black")


def make_main_flowchart():
    img = Image.new("RGB", (1800, 2500), "white")
    d = ImageDraw.Draw(img)
    title = "图 1  RoadGen3D 当前生成、修改与反馈主流程"
    tw = d.textbbox((0, 0), title, font=FONT_TITLE)[2]
    d.text(((1800 - tw) / 2, 35), title, fill="black", font=FONT_TITLE)
    note = "实线：系统数据流    虚线：人工反馈/重新生成路径"
    nw = d.textbbox((0, 0), note, font=FONT_SMALL)[2]
    d.text(((1800 - nw) / 2, 96), note, fill="black", font=FONT_SMALL)

    y_positions = [180, 475, 770, 1065, 1360, 1655, 1950]
    texts = [
        ("S1", "选取 AOI / OSM / GeoJSON\n标准化道路与周边要素\n输出：待审核二维参考标注"),
        ("S2", "人工审核二维事实\n保存新的不可变 source version\n记录：source_id / 父版本 / 指纹"),
        ("S3", "输入设计目标与权重\n归一化并生成确定性参数 patch\n默认：45 / 35 / 20，seed=42"),
        ("S4", "解析生成模式\nbaseline / auto → parametric\n仅显式 llm → 调用 LLM"),
        ("S5", "校验并合并参数\n白名单 + 数值边界 + 来源优先级\n+ 课程运行时固定约束"),
        ("S6", "约束求解与三维场景编排\n输出：scene_layout.json + scene.glb\n+ 参数来源与生产步骤"),
        ("S7", "评价、对比或局部 3D 编辑\n输出：指标 / score_delta / 新子版本\n只说明差异与相关，不说明因果"),
    ]
    for i, (sid, txt) in enumerate(texts):
        rect = (390, y_positions[i], 1410, y_positions[i] + 220)
        _box(d, rect, txt, title=sid)
        if i < len(texts) - 1:
            _arrow(d, (900, rect[3]), (900, y_positions[i+1]))

    _box(d, (430, 2265, 1370, 2430), "获得带参数来源、二维源版本、三维版本谱系和评价记录的\n可追踪三维街道场景", title="最终技术输出", width=7)
    _arrow(d, (900, 2170), (900, 2265))

    _box(d, (25, 1300, 315, 1585), "几何或背景建筑问题\n返回二维审核\n生成新 source", title="反馈 A", dashed=True, font=FONT_TINY)
    _arrow(d, (390, 1430), (315, 1430), dashed=True)
    _arrow(d, (170, 1300), (390, 585), dashed=True)

    _box(d, (1485, 1250, 1775, 1535), "目标或参数问题\n调整权重/建议\n重新生成", title="反馈 B", dashed=True, font=FONT_TINY)
    _arrow(d, (1410, 1380), (1485, 1380), dashed=True)
    _arrow(d, (1630, 1250), (1410, 880), dashed=True)

    _box(d, (1485, 1870, 1775, 2155), "局部资产问题\n受限 3D 命令\n仅创建子 revision", title="反馈 C", dashed=True, font=FONT_TINY)
    _arrow(d, (1410, 2060), (1485, 2060), dashed=True)
    _arrow(d, (1630, 2155), (1410, 2070), dashed=True)
    d.text((25, 2460), "注：当前实现不自动把 3D 修改回写为二维地理事实。", fill="black", font=FONT_TINY)
    png = FIG / "fig1_main_trace_flow.png"
    img.save(png)
    return png


def make_boundary_diagram():
    img = Image.new("RGB", (2200, 1250), "white")
    d = ImageDraw.Draw(img)
    title = "图 2  二维事实层、参数决策层与三维版本层的交互边界"
    tw = d.textbbox((0, 0), title, font=FONT_TITLE)[2]
    d.text(((2200 - tw) / 2, 35), title, fill="black", font=FONT_TITLE)
    _box(d, (60, 230, 520, 520), "Source v1\n待审核 annotation", title="二维事实层")
    _box(d, (650, 230, 1110, 520), "Source v2\napproved + sha256\n不可变", title="人工审核后")
    _arrow(d, (520, 375), (650, 375))
    _box(d, (650, 660, 1110, 930), "目标权重 / prompt\n规则映射或受限 LLM 建议\n用户接受字段", title="参数决策层", font=FONT_TINY)
    _arrow(d, (880, 520), (880, 660))
    _box(d, (1260, 230, 1650, 520), "Revision R1\nbaseline\nlayout + GLB", title="三维版本层")
    _box(d, (1740, 230, 2140, 520), "Revision R2\nredesign\n父版本=R1", title="重新生成")
    _arrow(d, (1110, 375), (1260, 375))
    _arrow(d, (1650, 375), (1740, 375))
    _arrow(d, (1110, 790), (1940, 520))
    _box(d, (1740, 660, 2140, 930), "Revision R3\n受限 3D 局部编辑\n父版本=R2", title="子版本", font=FONT_TINY)
    _arrow(d, (1940, 520), (1940, 660))
    _arrow(d, (1740, 790), (1110, 430), dashed=True)
    d.text((1340, 760), "不自动回写", fill="black", font=FONT_SMALL)
    _box(d, (250, 1020, 1950, 1190), "关键边界：parent_revision_id 记录版本谱系；当前 redesign 仍从已批准的二维 source 重新生成，\n不是在父 3D layout 上连续增量推演。若要改变道路几何或背景建筑 footprint，必须返回二维层。", font=FONT_TINY)
    png = FIG / "fig2_2d_3d_boundary.png"
    img.save(png)
    return png


MAIN_FLOW = make_main_flowchart()
BOUNDARY = make_boundary_diagram()


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="777777", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_column_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(h)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "E7ECEF")
        set_cell_margins(cell)
        if widths:
            set_column_width(cell, widths[i])
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(font_size)
                r.font.name = "Heiti SC"
    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = str(val)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if widths:
                set_column_width(cell, widths[i])
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                for r in p.runs:
                    r.font.size = Pt(font_size)
                    r.font.name = "Heiti SC"
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_field(run, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = text
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)
    return p


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.add_run(item)
    return


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_callout(doc, title, text, fill="EFF3F5"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="48545C", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    r = p.add_run(title + "　")
    r.bold = True
    p.add_run(text)
    return table


def set_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def build_docx():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.header_distance = Inches(0.35)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Heiti SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (("Title", 26, "1E2A32"), ("Heading 1", 17, "1E2A32"),
                              ("Heading 2", 13, "36464F"), ("Heading 3", 11, "36464F")):
        s = styles[name]
        s.font.name = "Heiti SC"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = True
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(5)

    for list_name in ("List Bullet", "List Bullet 2", "List Number"):
        styles[list_name].font.name = "Heiti SC"
        styles[list_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
        styles[list_name].font.size = Pt(10.2)
        styles[list_name].paragraph_format.space_after = Pt(2)

    # Header/footer
    hp = sec.header.paragraphs[0]
    hp.text = "ROADGEN3D｜技术交底书（当前实现快照）"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in hp.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(90, 100, 106)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("供发明人与专利代理师复核　·　")
    add_field(fp.add_run(), "PAGE")

    # Memo masthead
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run("技术交底书").bold = True
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor(67, 82, 90)
    title = doc.add_paragraph(style="Title")
    title.add_run("RoadGen3D 当前 LLM 参数、\n二维—三维交互与修改追踪")
    subtitle = doc.add_paragraph()
    subtitle.add_run("一种基于不可变二维参考标注和受限参数建议的可追踪道路三维场景生成与迭代方法").italic = True
    subtitle.paragraph_format.space_after = Pt(14)

    add_table(doc, ["文档属性", "当前值"], [
        ("代码快照", "RoadGen3D main｜d6f57cf2984ae65095dd0d75f3c24b7688e7aaeb｜2026-07-20 10:10 +08:00"),
        ("交底范围", "当前 LLM 模式与配置、参数决策、2D/3D 交互、修改依据、版本结果与评价闭环"),
        ("分析基准", "以当前源代码、当前非秘密环境配置和项目文档为准；未修改产品代码或运行配置"),
        ("复核状态", "供发明人与专利代理师复核；不是新颖性检索结论、可专利性意见或法律意见"),
    ], widths=[1.35, 5.05], font_size=9)

    doc.add_heading("阅读结论", level=1)
    add_callout(doc, "一句话结论", "系统当前不是“LLM 直接生成 3D”，而是以二维事实版本为输入、以确定性参数化生成器为主，LLM 仅在显式路径或专家参数建议界面中提出受限的参数修改；每次生成或 3D 编辑都形成可追踪的新版本。")
    add_bullets(doc, [
        "当前解析出的文本模型与视觉模型均为 gpt-4o-mini；协议为 OpenAI Chat Completions 兼容协议。",
        "课程端 baseline 和 redesign_default 均为 parametric；界面按钮虽然显示“让 LLM 设计下一版”，实际提交 auto，后端仍解析为 parametric。",
        "默认目标权重 45/35/20 经规则公式得到 sidewalk_width_m=3.22、density=0.96、seed=42 等参数。",
        "2D 审核生成不可变 source；3D 生成和局部编辑生成 revision。3D 不自动回写 2D，几何问题需人工返回 2D 重审。",
        "默认 structured 评价不调用视觉 LLM；版本比较只支持可追踪差异和相关性，不宣称因果。",
    ])

    doc.add_page_break()
    doc.add_heading("1. 交底目的与技术问题", level=1)
    doc.add_paragraph("本交底书回答五个面向实现的问题：当前调用何种 LLM、当前参数如何产生、参数在哪一步被修改、什么依据促成修改，以及修改后形成什么机器可读和可视化结果。")
    add_numbered(doc, [
        "如何把经人工确认的二维道路事实稳定地转换为三维街道场景，而不让生成过程篡改原始地理事实？",
        "如何让规则、用户输入和 LLM 同时参与参数决策，但仍能知道最终采用了哪一个值及其来源？",
        "如何将生成、人工三维编辑和评价结果组织成可回放、可比较的版本谱系？",
        "如何将结构化评价与视觉 LLM 评价分离，避免缺失或不稳定的视觉判断进入正式综合分？",
    ])
    doc.add_heading("1.1 系统定位", level=2)
    add_table(doc, ["层次", "当前实现", "不应误述为"], [
        ("二维事实层", "AOI/OSM/GeoJSON → 审核 → 不可变 source version", "可被 3D 编辑自动覆盖的二维图片"),
        ("参数决策层", "规则公式 + 显式输入 + 受限 LLM 建议 + 运行时固定约束", "LLM 自由改写道路拓扑"),
        ("三维编排层", "约束求解、资产放置、layout + GLB", "端到端 LLM 直接生成模型"),
        ("版本评价层", "revision 谱系、结构化指标、可选视觉评价", "无版本边界的覆盖保存"),
    ], widths=[1.15, 3.25, 2.0])
    doc.add_heading("1.2 建议技术名称", level=2)
    doc.add_paragraph("一种基于不可变二维参考标注和受限参数建议的可追踪道路三维场景生成与迭代方法。该名称强调“事实版本—参数约束—三维版本—评价记录”的组合关系，而非笼统强调大模型。")

    doc.add_page_break()
    doc.add_heading("2. 当前 LLM 模式、模型与运行参数", level=1)
    doc.add_heading("2.1 当前有效配置快照", level=2)
    add_table(doc, ["配置项", "当前有效值", "说明/来源"], [
        ("configured", "true", "已发现兼容端点和 API key；密钥未写入本文"),
        ("provider label", "openai", "ROADGEN_LLM_PROVIDER 未设置时的默认标签"),
        ("protocol", "openai_chat_completions", "兼容 OpenAI Chat Completions"),
        ("base URL", "https://api.zetatechs.com/v1/", "当前 .env 中 GRAPHRAG_API_BASE"),
        ("endpoint fingerprint", "sha256:86884b773945…f123f62d", "用于不泄露端点细节的能力指纹"),
        ("text model", "gpt-4o-mini", "未设置 ROADGEN_LLM_MODEL/LLM_MODEL，使用代码默认值"),
        ("vision model", "gpt-4o-mini", "未单独设置视觉模型，继承 text model"),
        ("temperature", "默认 0.2", "场景语义解析 0.1；视觉来源分类 0.0"),
        ("timeout", "120 s", "单次请求超时"),
        ("retry", "最多 10 次", "429 时从 4 s 起指数退避并加入抖动"),
        ("max_tokens / top_p / seed", "普通请求未设置", "由兼容端点/模型默认行为处理"),
    ], widths=[1.35, 1.85, 3.20], font_size=8.3)
    add_callout(doc, "验证边界", "本次只核验当前配置解析结果，没有向远程模型发送联网测试请求。因此“已配置”不等于本文已经证明远程端点此刻可达。", fill="F5F1E8")
    doc.add_heading("2.2 配置解析优先级", level=2)
    doc.add_paragraph("模型：ROADGEN_LLM_MODEL → LLM_MODEL → gpt-4o-mini；视觉模型：ROADGEN_LLM_VISION_MODEL → 文本模型。Base URL 与 API key 优先读取 ROADGEN_LLM_*，再读取 GRAPHRAG_*，然后读取 OPENAI_*，最后使用代码默认端点。")
    doc.add_paragraph("证据：C001、C002。")

    doc.add_page_break()
    doc.add_heading("3. 当前默认模式与界面—后端差异", level=1)
    add_callout(doc, "需要发明人确认的现状", "课程界面按 llm_configured 显示“已连接 LLM”和“让 LLM 设计下一版”，但按钮调用 generate(\"auto\")；当前后端把 auto 明确解析为 parametric。因此该按钮在当前默认课程路径中不会调用 LLM。", fill="F5EAEA")
    add_table(doc, ["入口/条件", "提交值", "后端解析", "是否调用 LLM", "结果"], [
        ("课程 baseline", "parametric", "parametric", "否", "规则/约束驱动基线版本"),
        ("课程 redesign 默认", "auto", "parametric", "否", "按新权重确定性重新生成"),
        ("显式 generation_mode=llm", "llm", "llm", "是，若 configured", "LLM 参数推导；失败会使该路径报错"),
        ("专家参数建议", "proposal API", "side-effect-free proposal", "是", "返回 before/after/reason/confidence，用户勾选后再以 parametric 生成"),
        ("LLM 未配置而请求 llm", "llm", "parametric fallback", "否", "记录 fallback_reason"),
    ], widths=[1.2, 1.0, 1.25, 1.0, 1.95], font_size=8.1)
    doc.add_heading("3.1 当前课程能力声明", level=2)
    add_bullets(doc, [
        "baseline = parametric；redesign_default = parametric；parametric_fallback = true。",
        "LLM parameter proposals 在配置有效时可用，但不是课程默认生成器。",
        "RAG mode = disabled，课程路径 knowledge_source = none。",
    ])
    doc.add_paragraph("证据：C003、C004、C006、C010、C011。")

    doc.add_page_break()
    doc.add_heading("4. 总体流程：参数在何时、为何被修改", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(MAIN_FLOW), width=Inches(5.15))
    doc.add_paragraph("图 1 将“修改的触发—依据—执行位置—输出结果”串成一条可审计链。三条虚线回路是不同含义的反馈：几何问题回到 2D；目标问题回到参数输入；局部资产问题只创建 3D 子版本。", style="Caption")

    doc.add_page_break()
    doc.add_heading("5. 二维与三维如何交互", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(BOUNDARY), width=Inches(6.25))
    doc.add_paragraph("图 2 突出了两个容易误解的边界：其一，3D 不自动反写 2D；其二，redesign 的 parent_revision_id 当前主要用于版本谱系，生成函数仍重新读取已批准的二维 source，而不是加载父 3D layout 做连续增量生成。", style="Caption")
    doc.add_heading("5.1 二维版本", level=2)
    add_bullets(doc, [
        "人工审核动作保存为新的 source version，记录 parent_source_id、review_actions、reviewed_at 和 annotation_sha256。",
        "后续生成通过 source_id 引用该二维事实输入，保留可回溯边界。",
    ])
    doc.add_heading("5.2 三维版本", level=2)
    add_bullets(doc, [
        "每个 SceneRevisionRecord 记录 source_id、parent_id、branch_kind、命令、scene_layout/GLB 路径及 provenance。",
        "局部 3D 编辑是受限命令并产生子 revision；锁定的 OSM/背景建筑不可编辑。",
        "道路几何或背景建筑 footprint 若需改变，必须返回 2D 审核并生成新的 source version。",
    ])

    doc.add_page_break()
    doc.add_heading("6. 当前默认参数：从目标权重到场景 patch", level=1)
    doc.add_paragraph("课程界面的默认目标权重是 walkability=45、safety=35、beauty=20。先归一化，再通过显式规则生成参数 patch。")
    add_equation(doc, "w_i = u_i / ∑_j u_j")
    add_equation(doc, "sidewalk_width_m = round(2.4 + 1.25w_walk + 0.65w_safe + 0.15w_beauty, 2)")
    add_equation(doc, "density = round(0.82 + 0.12w_walk + 0.05w_safe + 0.34w_beauty, 2)")
    add_table(doc, ["参数", "修改前/输入", "规则依据", "当前结果", "结果用途"], [
        ("权重", "45 / 35 / 20", "总和归一化", "0.45 / 0.35 / 0.20", "驱动规则映射与评价"),
        ("design_rule_profile", "—", "步行权重占优", "pedestrian_priority_v1", "选择设计规则"),
        ("objective_profile", "—", "当前权重组合", "balanced", "记录目标画像"),
        ("style_preset", "—", "步行+美观组合", "lush_walkable_v1", "风格选择"),
        ("sidewalk_width_m", "—", "公式 E002", "3.22 m", "人行道宽度"),
        ("density", "—", "公式 E003", "0.96", "街道家具全局密度"),
        ("ped_demand_level", "—", "步行权重阈值", "high", "行人需求"),
        ("bike_demand_level", "—", "安全/步行组合", "medium", "骑行需求"),
        ("furniture profile", "—", "课程规则", "pedestrian_friendly", "家具配置"),
        ("skeleton profile", "—", "课程规则", "walkable_commercial", "道路骨架配置"),
        ("seed", "—", "课程固定值", "42", "可重复性"),
    ], widths=[1.28, 1.0, 1.18, 1.28, 1.65], font_size=8)
    doc.add_paragraph("来源：C004、C005；公式：E001–E003。")

    doc.add_page_break()
    doc.add_heading("7. 参数修改的触发、依据、步骤与结果", level=1)
    add_table(doc, ["触发因素", "依据", "修改步骤", "被修改对象", "形成的结果"], [
        ("用户改变 45/35/20 目标权重", "归一化与确定性公式", "S3", "profile、人行道宽度、密度、需求等级、seed", "新的 compose patch；同输入可复现"),
        ("用户输入明确设计要求", "prompt_input / explicit_input", "S3–S5", "允许的场景参数", "较高来源优先级的候选值"),
        ("专家界面请求 LLM 建议", "提示词、参数规范、当前值、允许字段", "S4–S5", "skeleton、furniture、building representation、seed", "before/after/reason/confidence；未接受前无副作用"),
        ("用户勾选 LLM 建议", "逐字段人工确认", "S5", "所选参数", "field source=llm_suggestion；再以 parametric 生成"),
        ("课程运行时约束", "教学产品边界", "S5", "建筑表现、周边模式、land-use/infill、knowledge source 等", "覆盖较低优先级候选值"),
        ("生成器约束求解", "二维几何、参数边界、资产兼容性", "S6", "槽位、资产位置/方向/尺度", "scene_layout.json + scene.glb"),
        ("用户执行 3D 局部编辑", "命令白名单、支撑面和锁定校验", "S7", "可编辑资产变换/替换/增删", "新子 revision，不覆盖父版本"),
        ("评价或版本对比", "结构指标/可选视觉结果", "S7→S2/S3/S7", "二维事实、目标参数或局部资产", "建议、score_delta 与下一轮人工决策"),
    ], widths=[1.05, 1.25, 0.55, 1.65, 1.9], font_size=7.6)
    add_callout(doc, "审计链", "修改不能只记录“新值”；至少需要保留触发者、依据、所在步骤、修改前值、修改后值、字段来源、所属 source/revision 和输出 artifact。")

    doc.add_page_break()
    doc.add_heading("8. LLM 参数建议：允许范围、禁止范围与用户确认", level=1)
    doc.add_paragraph("当前较安全、最符合“AI 只做参数建议”定位的路径，是专家参数设计界面：LLM 返回最小化、无副作用的 patch；系统展示修改前/后、理由与置信度；用户逐项选择后，最终仍由确定性参数化生成器执行。")
    add_table(doc, ["类别", "LLM 可建议字段", "明确禁止"], [
        ("道路骨架", "laneCount、laneWidthM、sidewalkWidthM、furnishingWidthM、curbWidthM、junctionCornerPolicy/Radius", "道路中心线、拓扑、源 geometry"),
        ("街道家具", "globalDensity；分类 enabled、targetCountPer100M、spacing、setback、allowedZones", "资产 ID、直接场景编辑命令、GLB"),
        ("建筑", "representation", "building footprint、OSM 背景几何"),
        ("随机性", "seed", "不受控随机执行"),
        ("文件/事实", "无", "GeoJSON、approved annotation、source version"),
    ], widths=[1.0, 3.2, 2.2], font_size=8.2)
    doc.add_heading("8.1 参数规范中的数值边界", level=2)
    add_table(doc, ["字段", "允许范围/规则"], [
        ("laneCount", "整数 1–8"), ("laneWidthM", "2.5–4.5 m"), ("sidewalkWidthM", "1–12 m"),
        ("furnishingWidthM", "0–5 m"), ("curbWidthM", "0.05–0.4 m"),
        ("junctionCornerRadiusM", "fixed 策略下 1–20 m"), ("globalDensity", "0–2"),
        ("targetCountPer100M", "0–20"), ("preferred/minimumSpacingM", "2–100 m"),
        ("roadSetbackM", "0–10 m"), ("seed", "0–2,147,483,647"),
        ("锁定", "geometryLocked=true；footprintLocked=true"),
    ], widths=[2.2, 4.2], font_size=8.3)

    doc.add_page_break()
    doc.add_heading("9. 参数来源优先级与冲突解决", level=1)
    doc.add_paragraph("当同一字段出现多个候选值时，当前设计运行时按来源优先级选择。课程运行时固定 patch 在草案参数之后合并，因此能够覆盖规则或 LLM 候选。")
    add_table(doc, ["优先级", "来源", "含义"], [
        ("60", "explicit_input", "调用方明确输入"),
        ("57", "style blend / transfer", "风格融合或迁移"),
        ("55", "runtime_fixed", "课程/运行时不可变约束"),
        ("50", "scenario_hard_constraint", "场景硬约束"),
        ("45", "prompt_input", "自然语言中明确抽取的输入"),
        ("40", "parameter_triple", "结构化参数三元组"),
        ("30", "rag_supported_llm", "有检索证据支持的 LLM 候选；课程当前不启用"),
        ("20", "preset_default", "预设默认值"),
        ("15", "llm_derived", "LLM 推导候选"),
        ("0", "default_after_llm", "LLM 后兜底默认"),
    ], widths=[0.85, 2.25, 3.3], font_size=8.4)
    add_callout(doc, "课程当前固定覆盖", "building_representation=transparent_massing；surrounding_building_mode=footprint_based；auto_land_use_mode=off；infill_policy=off；building_height_mode=class_only；knowledge_source=none；retain_glb_policy=always；seed=42。")
    doc.add_paragraph("证据：C006、C012、C013。")

    doc.add_page_break()
    doc.add_heading("10. 三维编辑命令与安全边界", level=1)
    add_table(doc, ["编辑命令", "作用", "关键校验/结果"], [
        ("move / rotate", "移动或旋转可编辑资产", "不得命中锁定对象；写入新 revision"),
        ("scale", "尺度调整", "范围 0.25–4.0"),
        ("delete", "删除可编辑资产", "不允许删除 OSM 白模/背景上下文"),
        ("add / duplicate", "新增或复制资产", "资产与放置位置需合法"),
        ("replace_asset", "替换资产", "保留命令和 provenance"),
        ("set_building_style", "设置可编辑建筑表现", "footprint 锁定不因此解除"),
        ("auto_plant_trees", "自动植树", "校验种植带/家具带/临街面支撑区域"),
    ], widths=[1.4, 2.0, 3.0], font_size=8.4)
    doc.add_heading("10.1 course_grounded 协议", level=2)
    add_bullets(doc, [
        "高度偏移固定为 0；一般高度偏移边界为 0–10 m。",
        "树木允许落在 planting / furnishing / frontage；其他家具允许落在 sidewalk / furnishing / frontage。",
        "viewer manifest 中 editable=false 或 selection_source=osm_white_massing 的对象会被拒绝编辑。",
        "编辑成功后产生新的 scene_layout 和 GLB，并以 parent_id 连接父版本。",
    ])
    doc.add_paragraph("证据：C008、C009、P001。")

    doc.add_page_break()
    doc.add_heading("11. 生成结果、评价结果与证据产物", level=1)
    add_table(doc, ["产物", "内容", "可证明的事实"], [
        ("approved source record", "source_id、parent_source_id、annotation_sha256、review_actions", "使用了哪个二维事实版本"),
        ("scene_layout.json", "资产、变换、槽位、参数来源和布局事实", "机器可读场景结果"),
        ("scene.glb", "可视化三维模型", "所呈现的三维结果"),
        ("SceneRevisionRecord", "parent_id、branch_kind、commands、artifact、provenance", "生成/编辑形成何种版本关系"),
        ("parameter proposal", "before、after、reason、confidence、sideEffectFree", "LLM 建议了什么以及为什么"),
        ("structured evaluation", "可复现结构指标、建议与综合分", "不依赖视觉 LLM 的正式结构评价"),
        ("full visual evaluation", "基于渲染视图的可选 LLM 结果", "单独记录；不进入 formal structured composite"),
        ("revision comparison", "score_delta、差异记录、claim_scope", "版本差异和相关性，不证明因果"),
    ], widths=[1.45, 2.7, 2.25], font_size=8.1)
    doc.add_heading("11.1 评价模式", level=2)
    add_bullets(doc, [
        "默认 evaluation_mode=structured，不调用视觉 LLM。",
        "full 模式可在存在渲染视图时启用视觉 LLM。",
        "若视觉安全/美观分数不可用，则保持 N/A，不用虚构值填补。",
        "默认评价权重仍为 45/35/20，并在评价时归一化。",
    ])

    doc.add_page_break()
    doc.add_heading("12. 可支持的核心技术构思与技术效果", level=1)
    doc.add_paragraph("以下内容是基于当前实现整理的“技术构思候选”，用于发明人和代理师进一步提炼，不等于正式权利要求。")
    add_table(doc, ["特征组合", "技术作用", "可支持效果", "证据"], [
        ("不可变二维 source + 指纹 + 三维 revision 引用", "隔离地理事实与生成结果", "可回溯，避免 3D 编辑覆盖事实输入", "F001"),
        ("目标权重归一化 + 显式公式 patch", "把主观目标转换为可执行参数", "同输入与 seed 下可复现", "F002"),
        ("LLM 最小 patch + 白名单/边界 + 人工逐项接受", "限制 AI 副作用", "建议可拒绝、可审计、可回放", "F003"),
        ("二维几何 + 参数 patch + runtime_fixed → layout + GLB", "编译结构化意图", "保留机器可读中间结果和展示结果", "F004"),
        ("受限 3D 命令 + 支撑面校验 + 子版本", "控制局部修改", "减少无支撑摆放和地理上下文破坏", "F005"),
        ("structured/full 评价分离", "隔离可复现指标与视觉判断", "避免视觉缺失值冒充正式综合分", "F006"),
    ], widths=[2.05, 1.4, 2.35, 0.6], font_size=7.7)
    add_callout(doc, "建议核心构思", "以不可变二维事实版本为生成依据，将规则或经人工选择的受限 LLM 参数建议编译为带来源的参数规范，再通过约束求解形成三维场景布局和可视化模型，并以不可变三维子版本和分离式评价记录修改结果。")

    doc.add_page_break()
    doc.add_heading("13. 当前实现缺口与不宜过度主张的内容", level=1)
    add_table(doc, ["当前缺口/边界", "交底书中的处理", "后续确认"], [
        ("课程按钮文案暗示 LLM，但 auto 实际为 parametric", "明确列为界面—后端差异，不把默认路径写成 LLM 生成", "是否改文案或改调用模式"),
        ("parent_revision_id 不等于加载父 3D layout 增量生成", "只主张版本谱系，不主张连续三维推演", "是否计划实现增量 redesign"),
        ("3D 不自动回写 2D", "将回到二维定义为人工反馈环", "是否需要设计显式反向提议协议"),
        ("课程 RAG disabled / knowledge_source=none", "不主张当前课程有知识检索证据闭环", "未来启用时需保存引用与版本"),
        ("structured 评价不含视觉 LLM", "分离陈述，不把视觉结果写入正式综合分", "代理师判断是否作为从属构思"),
        ("未开展现有技术检索", "只整理实现证据，不作新颖性/创造性结论", "由代理师安排检索"),
    ], widths=[2.35, 2.5, 1.55], font_size=8.1)
    doc.add_heading("13.1 发明人需要补充回答的问题", level=2)
    add_numbered(doc, [
        "最希望保护的是“受限 LLM 参数建议”本身，还是“二维事实—参数—三维版本—评价”的整体闭环？",
        "当前 45/35/20 和两条映射公式是教学示例、优选实施例，还是必须保留的核心技术特征？",
        "是否已有实际测试数据证明可重复性、编辑合法率、人工用时或评价稳定性的提升？",
        "未来是否会让 3D 编辑生成二维变更提议？如会，需要定义审核和冲突解决机制。",
        "Zeta 兼容端点属于部署示例还是固定供应商？建议权利化表述保持协议/能力层中立。",
    ])

    doc.add_page_break()
    doc.add_heading("14. 实施方式示例：一次完整修改如何被记录", level=1)
    doc.add_paragraph("以下用当前默认值说明一条可重放的实施链。")
    add_table(doc, ["阶段", "输入/状态", "动作及依据", "输出/记录"], [
        ("S1–S2", "OSM/GeoJSON + 人工审核", "确认道路和背景事实；保存新 source", "source_id、parent_source_id、annotation_sha256"),
        ("S3", "45/35/20", "归一化为 0.45/0.35/0.20；应用 E002/E003", "sidewalk=3.22m；density=0.96；profiles；seed=42"),
        ("S4", "generation_mode=auto", "按当前课程策略解析", "resolved_generation_mode=parametric"),
        ("S5", "规则 patch + 固定约束", "校验范围并以 runtime_fixed 覆盖冲突", "transparent_massing、knowledge_source=none 等有效参数"),
        ("S6", "approved annotation + final parameters", "约束求解、槽位和资产编排", "scene_layout.json、scene.glb、production_steps"),
        ("S7", "结构化评价 45/35/20", "计算指标和建议；比较父子版本", "evaluation、score_delta、非因果声明"),
        ("可选编辑", "用户移动一项可编辑家具", "通过锁定/支撑面/范围校验", "新 child revision；父版本不变；2D 不变"),
    ], widths=[0.72, 1.65, 2.15, 1.88], font_size=8.1)
    add_callout(doc, "修改依据的落点", "参数修改理由应落在 proposal.reason、规则公式或硬约束记录中；修改结果应同时落在 final parameter source、scene_layout、revision provenance 和评价记录中，而不能只存在于界面提示文字。")

    doc.add_page_break()
    doc.add_heading("附录 A　源代码与文档证据索引", level=1)
    rows = []
    for item in SOURCES["source_map"]:
        rows.append((item["id"], item["type"], item["locator"], item["summary"], item["confidence"]))
    add_table(doc, ["ID", "类型", "定位", "支持内容", "置信"], rows,
              widths=[0.55, 0.9, 2.15, 2.35, 0.45], font_size=6.8)

    doc.add_page_break()
    doc.add_heading("附录 B　术语与证据边界", level=1)
    term_rows = []
    for item in INVENTORY["terminology_ledger"]:
        term_rows.append((item["canonical_zh"], "、".join(item["source_terms"]), "、".join(item["forbidden_aliases"])))
    add_table(doc, ["本文术语", "源代码术语", "避免的误称"], term_rows,
              widths=[1.75, 2.6, 2.05], font_size=8.2)
    doc.add_heading("证据状态说明", level=2)
    add_bullets(doc, [
        "explicit：当前代码或文档直接支持。本文核心技术事实 F001–F007 均为 explicit。",
        "当前环境模型值来自代码解析和非秘密环境字段；密钥未输出。",
        "技术效果按机制可推导，但若需要量化效果，仍须补充实验、对照、日志或用户测试数据。",
        "本交底书没有进行专利数据库检索，不能替代代理师的新颖性和创造性判断。",
    ])
    doc.add_paragraph("文件结束。", style="Subtitle")

    # Global paragraph and table safeguards.
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            p.paragraph_format.keep_with_next = True
        for r in p.runs:
            if not r.font.name:
                r.font.name = "Heiti SC"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Heiti SC")
    for table in doc.tables:
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.widow_control = True

    props = doc.core_properties
    props.title = "RoadGen3D 当前 LLM 参数、二维—三维交互与修改追踪技术交底书"
    props.subject = "RoadGen3D 当前实现证据化技术交底"
    props.author = "RoadGen3D 项目组（由 Codex 基于当前实现整理）"
    props.keywords = "RoadGen3D, LLM, 参数建议, 2D, 3D, 技术交底书, 版本谱系"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
